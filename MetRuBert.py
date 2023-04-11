import glob
import sys
from MetRobert_rel import main_dutch
# import our own output prettifier
import MetRubert_parser as outputgen
import os
import re

# import used for docx generation
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
inputfile = "dev.tsv"
main_dutch.main(inputfile)

# settings and modules
# whether to include this pos tag
noun = "yes"
verb = "yes"
adj = "yes"
adv = "yes"
pron = "yes"
det = "no"
num = "no"
# wheter to include softmax data or not
softmax = "yes"

print("What mode would you like to use? Type txt for simple text input and dev for dev.tsv input.")
mode = input("Please enter your choice: ")
while mode != "tok" and mode != "txt" and mode != "dev":
    mode = input("Invalid choice, try again: ")


def findReplace(directory, filePattern):
        for path, dirs, files in os.walk(os.path.abspath(directory)):
            for filename in fnmatch.filter(files, filePattern):
                filepath = os.path.join(path, filename)
                s = pathlib.Path(filepath).read_text()

                if s[:2] == "' ":
                    s = s.replace("' ", "'", 1)

                # extra checks added
                s = s.replace("''", "")
                s = s.replace(" ,,", " ")
                s = s.replace(",,", " ")
                s = s.replace("\"", "")
                s = s.replace(" , ", ", ")
                s = s.replace(" \\ ", "\\")
                s = s.replace(" / ", "/")
                s = s.replace(" :", ":")
                s = s.replace(",, ", "")

                # second round of extra checks
                s = s.replace(", ,", "")
                s = s.replace(" ?", "?")
                s = s.replace(" - ", "- ")

                # third round of extra checks
                s = s.replace(" ( ", " (")
                s = s.replace(" ) ", ") ")
                s = s.replace(" . ", ". ")

                s = s.replace(" ',", "',")
                s = s.replace(" '',", "'',")

                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                with open(filepath, "w") as f:
                    f.write(s)

def wordfunction(outputdir):

    # Create a document, add heading
    document = Document()
    document.add_heading('MetRuBert Output', 0)

    ######################
    # SENTENCE RETRIEVAL #
    ######################

    previous_sentence = None
    pos_list = []
    met_list = []

    sentences = []

    with open(os.path.join(outputdir, 'output.tsv'), mode='r', encoding='utf-8') as output:
        for line in output:
            if line[:5] != "index":
                splitted = line.split("\t")

                sentence = splitted[1]

                if previous_sentence == None or sentence == previous_sentence:
                    # append pos tag
                    pos = splitted[2]
                    index = splitted[3]
                    pos_list.append([index, pos])

                    # append metaphor
                    met = splitted[5]
                    if met.find("1") != -1 and met.find("-1") == -1:
                        met_list.append(index)
                elif sentence != previous_sentence:
                    split_sen = previous_sentence.split()
                    p = document.add_paragraph('')

                    for i in range(len(split_sen)):
                        word = split_sen[i]
                        found = False
                        if f'{str(i)} ' in met_list:
                            for pos in pos_list:
                                # print(pos)
                                if pos[0] == f'{str(i)} ':
                                    found = True
                                    if pos[1] == "NOUN":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.YELLOW
                                        font.bold = True
                                    elif pos[1] == "VERB":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                                        font.bold = True
                                    elif pos[1] == "DET":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                        font.bold = True
                                    elif pos[1] == "ADV":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.PINK
                                        font.bold = True
                                    elif pos[1] == "ADJ":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.RED
                                        font.bold = True
                                    elif pos[1] == "PRON":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                                        font.bold = True
                                    else:
                                        p.add_run(word).bold = True
                            if found == False:
                                p.add_run(word).bold = True
                        else:
                            for pos in pos_list:
                                # print(pos)
                                if pos[0] == f'{str(i)} ':
                                    found = True
                                    if pos[1] == "NOUN":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    elif pos[1] == "VERB":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                                    elif pos[1] == "DET":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.GRAY_25
                                    elif pos[1] == "ADV":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.PINK
                                    elif pos[1] == "ADJ":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.RED
                                    elif pos[1] == "PRON":
                                        font = p.add_run(word).font
                                        font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                                    else:
                                        print(word)
                                        print(pos[1])
                                        print(pos[0])
                                        print(pos_list)

                                        p.add_run(word)
                            if found == False:
                                # print(word)
                                # print(i)
                                # print(pos_list)
                                # Missing elements in list
                                p.add_run(word)

                        p.add_run(' ')
                    sentences.append(previous_sentence)

                    pos_list = []
                    met_list = []

                previous_sentence = sentence

    document.add_page_break()
    document.save(outputdir + '/' + 'MetRubert_runx.docx')

    def findReplace(directory, filePattern):
        for path, dirs, files in os.walk(os.path.abspath(directory)):
            for filename in fnmatch.filter(files, filePattern):
                filepath = os.path.join(path, filename)
                s = pathlib.Path(filepath).read_text()

                if s[:2] == "' ":
                    s = s.replace("' ", "'", 1)

                # extra checks added
                s = s.replace("''", "")
                s = s.replace(" ,,", " ")
                s = s.replace(",,", " ")
                s = s.replace("\"", "")
                s = s.replace(" , ", ", ")
                s = s.replace(" \\ ", "\\")
                s = s.replace(" / ", "/")
                s = s.replace(" :", ":")
                s = s.replace(",, ", "")

                # second round of extra checks
                s = s.replace(", ,", "")
                s = s.replace(" ?", "?")
                s = s.replace(" - ", "- ")

                # third round of extra checks
                s = s.replace(" ( ", " (")
                s = s.replace(" ) ", ") ")
                s = s.replace(" . ", ". ")

                s = s.replace(" ',", "',")
                s = s.replace(" '',", "'',")

                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                s = s.replace(" ' ", " '", 1)
                s = s.replace(" ' ", "' ", 1)
                with open(filepath, "w") as f:
                    f.write(s)

def shellsafe(s, quote="'"):
    """Return a shell-safe version of the input string."""
    # Escape any occurrences of the quote character
    s = s.replace(quote, "\\" + quote)

    # Quote the string if it contains whitespace or shell metacharacters
    if re.search(r'[\s\'"$`\\]', s):
        s = quote + s + quote

    return s

# use parameter values to create a parameter list
pos_list = []
if str(noun) == "no":
    pos_list.append("noun")
if str(verb) == "no":
    pos_list.append("verb")
if str(adj) == "no":
    pos_list.append("adj")
if str(adv) == "no":
    pos_list.append("adv")
if str(pron) == "no":
    pos_list.append("pron")
if str(det) == "no":
    pos_list.append("det")
if str(num) == "no":
    pos_list.append("num")

outputdir = os.path.realpath(
    os.path.join(os.getcwd(), os.path.dirname(__file__)))

if mode == "dev":
    dev_path = input("Please enter the path of the dev.tsv file")
elif mode == "txt":
    tok_folder = input(
        "Please enter the path of the folder containing the txt files")
    # get files in folder, filter on .tok files, retry if not present
    tok_files = [f for f in os.listdir(tok_folder) if os.path.isfile(
        os.path.join(tok_folder, f)) and f.endswith(".txt")]
    if len(tok_files) == 0:
        print("No txt files found in this folder")
        exit()

    # print the files in the folder and ask the user to select one (DEBUG)
    print("Found the following txt files in the folder:")
    for i in range(len(tok_files)):
        print(str(i) + " : " + tok_files[i])

    # iterate over files
    for i in range(len(tok_files)):

        print("Tokenizing file " + tok_files[i] + " ...")
        # ? 
        # ? RUNNING UCTO
        # ? 

        # get the path of the file
        tok_path = os.path.join(tok_folder, tok_files[i])
        inputfilepath = str(tok_path)
        basename = os.path.basename(inputfilepath)[:-4] # remove .txt

        # output tok path
        tokfile = os.path.join(outputdir, basename + '.tok')
        r = os.system('ucto -L nl -n ' + shellsafe(inputfilepath,
                        '"') + ' > ' + shellsafe(tokfile, '"'))
        if r != 0:
            print("Failure running ucto", file=sys.stderr)
            sys.exit(2)

        # Replace all quotation marks in files
        findReplace(outputdir, "*.txt")
        findReplace(outputdir, "*.tok")

        print("Parsing file with alpino: " + tok_files[i] + " ...")
        # ? 
        # ? RUNNING ALPINO
        # ?

        pwd = os.getcwd()
        os.chdir(outputdir)
        if not os.path.exists("xml"):
            os.mkdir("xml")
        else:
            for filename in glob.glob('xml/*.xml'):
                os.unlink(filename)  # clear for next round

        ALPINO_HOME = os.getenv('ALPINO_HOME')

        if ALPINO_HOME is None:
            print("ALPINO_HOME not set", file=sys.stderr)
            sys.exit(1)

        cmd = "ALPINO_HOME=" + shellsafe(ALPINO_HOME) + " " + ALPINO_HOME + \
            "/bin/Alpino -veryfast -flag treebank xml debug=1 end_hook=xml user_max=900000 -parse < " + tokfile
        print(cmd, file=sys.stderr)
        r = os.system(cmd)
        if r != 0:
            print("Failure running alpino", file=sys.stderr)
            sys.exit(2)

        os.chdir("xml")

        os.chdir('..')
        os.rename('xml', 'xml_' + basename)
        os.chdir(pwd)

    print("Generating dev.tsv file")
    #? 
    #? Generating dev.tsv
    #?    

    # get location needed
    dev_location = outputdir

    import pasmaparser_cov_melBert_allpos_clam as parser

    # run the python file to generate dev data
    parser.main(dev_location)

    # go to directory where dev.tsv data was created
    os.chdir(dev_location)

    #?
    #? Cleaning up intermediate files
    #?

    # cleanup folders unneeded xml files
    for dire in os.listdir(dev_location):
        d = os.path.join(dev_location, dire)
        if os.path.isdir(d):
            print("cleaning folder" + str(d))
            for file in os.listdir(d):
                if file.endswith(".xml") and str(alp) == "no":
                    try:
                        os.remove(d + "/" + file)
                    except:
                        print("Error while deleting xml file : ", file)

    # get location of dev file to copy
    dev_path = dev_location + "/" + "dev.tsv"

# TODO: use output folder (replace first none with outputdir)
outputgen.main(None, pos_list, False, dev_path, softmax)

if True:
    # cleanup
    for file in os.listdir(outputdir):
        if file.endswith("dev_float.txt"):  # and str(unres) == "no":
            try:
                os.remove(outputdir + "/" + file)
            except:
                print("Error while deleting tok file : ", file)
        elif file.endswith("dev_soft.txt"):  # and str(sof) == "no":
            try:
                os.remove(outputdir + "/" + file)
            except:
                print("Error while deleting tok file : ", file)
        elif file.endswith("_dev2.txt") or file.endswith("dev2.tsv") or file.endswith("soft2.txt"):
            try:
                os.remove(outputdir + "/" + file)
            except:
                print("Error while deleting tok file : ", file)

    # Generate word output
    wordfunction(outputdir=outputdir)
